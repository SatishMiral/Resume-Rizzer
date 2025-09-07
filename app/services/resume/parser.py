from docx import Document
import re
import nltk

# Using NLTK in Production can be triky
def extract_sentences(docx_path: str):
    doc = Document(docx_path)
    sentences = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  
            para_sentences = nltk.sent_tokenize(text)
            sentences.extend(para_sentences)
    return sentences

# Simpler version to extract sentences
def extract_sentences_regex(docx_path: str):
    doc = Document(docx_path)
    sentences = []
    id_counter = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # ignore empty lines
            # Split into sentences
            para_sentences = re.split(r'(?<=[.!?]) +', text)
            for s in para_sentences:
                s = s.strip()
                if s:
                    sentences.append({"id": id_counter, "text": s})
                    id_counter += 1

    return {"sentences": sentences}
