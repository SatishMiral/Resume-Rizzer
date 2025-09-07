from docx import Document

def replace_and_style(doc, from_sentence, to_sentence, bold_words=None, italic_words=None):
    bold_words = bold_words or []
    italic_words = italic_words or []

    for para in doc.paragraphs:
        if from_sentence in para.text:
            # Clear the paragraph text (remove old runs)
            for run in para.runs:
                run.text = ""

            # Add styled replacement text word by word
            for word in to_sentence.split(" "):
                run = para.add_run(word + " ")
                clean_word = word.strip(",.!?;:")

                if clean_word in bold_words:
                    run.bold = True
                if clean_word in italic_words:
                    run.italic = True

    # Also check inside tables 
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if from_sentence in cell.text:
                    # Clear and rebuild cell
                    cell.text = ""
                    para = cell.add_paragraph()
                    for word in to_sentence.split(" "):
                        run = para.add_run(word + " ")
                        clean_word = word.strip(",.!?;:")
                        if clean_word in bold_words:
                            run.bold = True
                        if clean_word in italic_words:
                            run.italic = True